"""
Resume Controller

Business logic for resume operations including submission,
retrieval, and AI analysis.

RBAC Rules:
- Applicants can submit resumes (submit_resume permission)
- Applicants can view only their own resumes (view_own_resume permission)
- Recruiters/Admins can view all resumes (view_all_resumes permission)
- Recruiters/Admins can trigger AI analysis (analyze_resume permission)
"""

import io
from typing import List
from sqlmodel import Session, select
from PyPDF2 import PdfReader
from docx import Document

from models.user import User
from models.resume import Resume
from schemas.resume import (
    ResumeSubmitRequest, 
    ResumeResponse, 
    ResumeListResponse,
    ResumeAnalysisResponse,
    AnalyzeResumeRequest
)
from services.gemini_service import analyze_resume as gemini_analyze, extract_resume_metadata
import json
from services.anonymization_service import anonymize_resume_content


def submit_resume(
    session: Session, 
    current_user: User, 
    request: ResumeSubmitRequest
) -> ResumeResponse:
    """
    Submit a new resume for the current user.
    
    This function creates a new resume record associated with
    the authenticated user. Only Applicants can submit resumes.
    
    RBAC Requirement:
        Permission: 'submit_resume' (Applicant)
    
    Args:
        session: Database session
        current_user: The authenticated user (must be Applicant)
        request: Resume content to submit
    
    Returns:
        ResumeResponse: The created resume data
    """
    # Call Gemini for one-time extraction
    metadata = extract_resume_metadata(request.content)
    
    resume = Resume(
        user_id=current_user.id,
        content=request.content,
        analysis_result=None,
        extracted_skills=json.dumps(metadata.get("skills", [])),
        experience_years=metadata.get("experience_years", 0)
    )
    
    session.add(resume)
    session.commit()
    session.refresh(resume)
    
    return ResumeResponse(
        id=resume.id,
        user_id=resume.user_id,
        username=current_user.username,
        content=resume.content,
        analysis_result=resume.analysis_result,
        created_at=resume.created_at
    )


def get_resumes(
    session: Session, 
    current_user: User,
    skip: int = 0,
    limit: int = 50
) -> tuple[ResumeListResponse, int, int]:
    """
    Retrieve resumes based on user permissions.
    
    RBAC Logic:
        - If user has 'view_all_resumes' permission (Recruiter/Admin):
          Returns ALL resumes in the system
        - If user has only 'view_own_resume' permission (Applicant):
          Returns only resumes owned by the user
    
    Args:
        session: Database session
        current_user: The authenticated user
    
    Returns:
        ResumeListResponse: List of accessible resumes
    """
    # Check if user can view all resumes (Recruiter/Admin)
    if current_user.has_permission("view_all_resumes"):
        # Query all resumes with user info
        statement = select(Resume)
    else:
        # Applicant: Only their own resumes
        statement = select(Resume).where(Resume.user_id == current_user.id)
        
    statement = statement.order_by(Resume.created_at.desc())
    
    # Pagination
    total = len(session.exec(statement).all())
    total_pages = (total + limit - 1) // limit if limit > 0 else 1
    
    resumes = session.exec(statement.offset(skip).limit(limit)).all()
    
    # Convert to response format
    resume_list = []
    for resume in resumes:
        # Get the owner's username
        owner = session.get(User, resume.user_id)
        resume_list.append(ResumeResponse(
            id=resume.id,
            user_id=resume.user_id,
            username=owner.username if owner else "Unknown",
            content=resume.content,
            analysis_result=resume.analysis_result,
            created_at=resume.created_at
        ))
    
    return ResumeListResponse(
        resumes=resume_list,
        total=total,
        page=(skip // limit) + 1 if limit > 0 else 1,
        limit=limit,
        total_pages=total_pages
    ), total, total_pages


def get_resume_by_id(session: Session, resume_id: int, current_user: User) -> ResumeResponse:
    """
    Retrieve a specific resume by ID.
    
    RBAC Logic:
        - Recruiters/Admins can view any resume
        - Applicants can only view their own resumes
    
    Args:
        session: Database session
        resume_id: ID of the resume to retrieve
        current_user: The authenticated user
    
    Returns:
        ResumeResponse: The requested resume
    
    Raises:
        ValueError: If resume not found or access denied
    """
    resume = session.get(Resume, resume_id)
    
    if not resume:
        raise ValueError(f"Resume with ID {resume_id} not found")
    
    # Check access permissions
    # Recruiters/Admins can view all, Applicants only their own
    if not current_user.has_permission("view_all_resumes"):
        if resume.user_id != current_user.id:
            raise ValueError("Access denied: You can only view your own resumes")
    
    owner = session.get(User, resume.user_id)
    
    return ResumeResponse(
        id=resume.id,
        user_id=resume.user_id,
        username=owner.username if owner else "Unknown",
        content=resume.content,
        analysis_result=resume.analysis_result,
        created_at=resume.created_at
    )


def analyze_resume(
    session: Session, 
    resume_id: int, 
    current_user: User,
    request: AnalyzeResumeRequest
) -> ResumeAnalysisResponse:
    """
    Trigger AI analysis on a resume using the Gemini API.
    
    This function:
    1. Retrieves the resume
    2. Sends content to Gemini API for analysis
    3. Stores the analysis result in the database
    4. Returns the updated resume with analysis
    
    RBAC Requirement:
        Permission: 'analyze_resume' (Recruiter/Admin only)
    
    Check if user role includes 'analyze_resume' permission before 
    calling Gemini. This is enforced at the route level.
    
    Args:
        session: Database session
        resume_id: ID of the resume to analyze
        current_user: The authenticated user (must have analyze_resume permission)
        request: Analysis parameters (job_role, additional_context)
    
    Returns:
        ResumeAnalysisResponse: Resume with AI analysis result
    
    Raises:
        ValueError: If resume not found
    """
    # Get the resume
    resume = session.get(Resume, resume_id)
    
    if not resume:
        raise ValueError(f"Resume with ID {resume_id} not found")
    
    # Anonymize PII before sending to AI
    # The original resume.content in the database stays untouched
    anonymized_content = anonymize_resume_content(resume.content)
    
    # Call Gemini API for analysis with anonymized content
    # The analyze_resume permission check happens at the route level
    analysis_result = gemini_analyze(
        content=anonymized_content,
        job_role=request.job_role,
        additional_context=request.additional_context
    )
    
    # Update the resume with analysis result
    resume.analysis_result = analysis_result
    session.add(resume)
    session.commit()
    session.refresh(resume)
    
    return ResumeAnalysisResponse(
        id=resume.id,
        content=resume.content,
        analysis_result=resume.analysis_result
    )


def delete_resume(session: Session, resume_id: int, current_user: User) -> bool:
    """
    Delete a resume.
    
    RBAC Logic:
        - Applicants can delete only their own resumes
        - Admins can delete any resume
    
    Args:
        session: Database session
        resume_id: ID of the resume to delete
        current_user: The authenticated user
    
    Returns:
        bool: True if deleted successfully
    
    Raises:
        ValueError: If resume not found or access denied
    """
    resume = session.get(Resume, resume_id)
    
    if not resume:
        raise ValueError(f"Resume with ID {resume_id} not found")
    
    # Check access - owner or admin can delete
    if resume.user_id != current_user.id:
        if not current_user.has_permission("manage_users"):  # Admin permission
            raise ValueError("Access denied: You can only delete your own resumes")
    
    session.delete(resume)
    session.commit()
    
    return True


# =============================================================================
# File Processing Functions
# =============================================================================

def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        file_content: Raw bytes of the PDF file
    
    Returns:
        str: Extracted text content
    
    Raises:
        ValueError: If PDF cannot be read or has no text
    """
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        if not text_parts:
            raise ValueError("PDF contains no extractable text")
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to read PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        file_content: Raw bytes of the DOCX file
    
    Returns:
        str: Extracted text content
    
    Raises:
        ValueError: If DOCX cannot be read or has no text
    """
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        if not text_parts:
            raise ValueError("DOCX contains no extractable text")
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {str(e)}")


def submit_resume_file(
    session: Session,
    current_user: User,
    file_content: bytes,
    filename: str,
    file_type: str
) -> ResumeResponse:
    """
    Submit a resume from an uploaded file (PDF or DOCX).
    
    This function:
    1. Extracts text content from the file
    2. Creates a resume record with the extracted text
    3. Stores file metadata for reference
    
    RBAC Requirement:
        Permission: 'submit_resume' (Applicant)
    
    Args:
        session: Database session
        current_user: The authenticated user (must be Applicant)
        file_content: Raw bytes of the uploaded file
        filename: Original filename
        file_type: File extension ('pdf' or 'docx')
    
    Returns:
        ResumeResponse: The created resume data
    
    Raises:
        ValueError: If file type is unsupported or text extraction fails
    """
    # Extract text based on file type
    if file_type == "pdf":
        content = extract_text_from_pdf(file_content)
    elif file_type == "docx":
        content = extract_text_from_docx(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    
    # Validate minimum content length
    if len(content) < 50:
        raise ValueError("Extracted text is too short (minimum 50 characters required)")
    
    # Call Gemini for one-time extraction
    metadata = extract_resume_metadata(content)
    
    # Create new resume linked to current user
    resume = Resume(
        user_id=current_user.id,
        content=content,
        analysis_result=None,
        original_filename=filename,
        file_type=file_type,
        extracted_skills=json.dumps(metadata.get("skills", [])),
        experience_years=metadata.get("experience_years", 0)
    )
    
    session.add(resume)
    session.commit()
    session.refresh(resume)
    
    return ResumeResponse(
        id=resume.id,
        user_id=resume.user_id,
        username=current_user.username,
        content=resume.content,
        analysis_result=resume.analysis_result,
        created_at=resume.created_at
    )
