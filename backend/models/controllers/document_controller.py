import io
import os
import uuid
from datetime import datetime
from typing import List, Optional
from sqlmodel import Session, select
from fastapi import UploadFile, HTTPException, status
from fastapi.responses import FileResponse
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from models.document import Document
from models.audit_log import AuditLog
from models.user import User
from services.automation_flags import is_automation_enabled
from services.automation_heuristics import extract_document_metadata

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "uploads")

# Ensure upload directory exists
os.makedirs(UPLOAD_DIR, exist_ok=True)

def _log_audit(session: Session, user_id: int, action: str, entity_type: str, entity_id: int, details: str = None):
    audit = AuditLog(
        user_id=user_id,
        action=action,
        entity_type=entity_type,
        entity_id=entity_id,
        details=details
    )
    session.add(audit)
    # We do not commit here because the parent transaction should commit


def _extract_text_from_file_bytes(file_bytes: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            return "\n".join(filter(None, (page.extract_text() or "" for page in reader.pages)))
        except Exception:
            return ""
    if ext == "docx":
        try:
            document = DocxDocument(io.BytesIO(file_bytes))
            return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)
        except Exception:
            return ""
    return file_bytes[:4000].decode("utf-8", errors="ignore")

def submit_document_file(
    session: Session,
    current_user: User,
    file: UploadFile,
    document_type: str,
    expiration_date: Optional[datetime] = None
) -> Document:
    # Validate file
    if not file.filename:
        raise ValueError("No filename provided")
        
    file_bytes = file.file.read()
    file_size = len(file_bytes)
    extracted_text = _extract_text_from_file_bytes(file_bytes, file.filename)
    metadata_hint = (
        extract_document_metadata(file.filename, extracted_text, document_type)
        if is_automation_enabled("doc_ocr_extract")
        else {
            "document_type_candidate": document_type,
            "expiry_date_candidate": expiration_date.isoformat() if expiration_date else None,
            "confidence": 0.0,
        }
    )
    
    # Generate unique storage filename
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    safe_filename = f"{uuid.uuid4().hex}.{ext}"
    file_path = os.path.join(UPLOAD_DIR, safe_filename)
    
    # Save physical file
    with open(file_path, "wb") as f:
        f.write(file_bytes)
        
    # Save DB record
    doc = Document(
        user_id=current_user.id,
        document_type=document_type,
        file_path=file_path,
        original_filename=file.filename,
        file_size_bytes=file_size,
        expiration_date=expiration_date,
        document_type_candidate=metadata_hint["document_type_candidate"],
        expiration_date_candidate=(
            datetime.fromisoformat(metadata_hint["expiry_date_candidate"])
            if metadata_hint["expiry_date_candidate"]
            else expiration_date
        ),
        extraction_confidence=metadata_hint["confidence"],
        metadata_confirmed=False,
    )
    
    session.add(doc)
    session.commit()
    session.refresh(doc)
    
    # Audit log
    _log_audit(
        session=session,
        user_id=current_user.id,
        action="UPLOAD_DOCUMENT",
        entity_type="Document",
        entity_id=doc.id,
        details=f"Uploaded {document_type} file: {file.filename}"
    )
    session.commit()
    
    return doc


def confirm_document_metadata(
    session: Session,
    doc_id: int,
    current_user: User,
    document_type: Optional[str] = None,
    expiration_date: Optional[datetime] = None,
) -> Document:
    doc = session.get(Document, doc_id)
    if not doc:
        raise ValueError("Document not found")

    if doc.user_id != current_user.id and not current_user.has_permission("manage_users"):
        raise ValueError("You do not have permission to update this document")

    before = {
        "document_type": doc.document_type,
        "expiration_date": doc.expiration_date.isoformat() if doc.expiration_date else None,
    }

    if document_type:
        doc.document_type = document_type
    if expiration_date is not None:
        doc.expiration_date = expiration_date

    doc.metadata_confirmed = True
    session.add(doc)
    session.commit()
    session.refresh(doc)

    _log_audit(
        session=session,
        user_id=current_user.id,
        action="CONFIRM_DOCUMENT_METADATA",
        entity_type="Document",
        entity_id=doc.id,
        details=f"Confirmed metadata for {doc.original_filename}",
    )
    session.commit()
    return doc

def get_documents(session: Session, current_user: User) -> List[Document]:
    # If admin/recruiter, return all documents. Otherwise just user's.
    if current_user.has_permission("view_all_resumes"):
        return session.exec(select(Document).order_by(Document.uploaded_at.desc())).all()
    else:
        return session.exec(
            select(Document).where(Document.user_id == current_user.id).order_by(Document.uploaded_at.desc())
        ).all()

def get_user_documents(session: Session, target_user_id: int, current_user: User) -> List[Document]:
    # Ensure they have permission to view other users' docs
    if not current_user.has_permission("view_all_resumes"):
        raise ValueError("You do not have permission to view this user's documents")
    
    return session.exec(
        select(Document)
        .where(Document.user_id == target_user_id)
        .order_by(Document.uploaded_at.desc())
    ).all()

def download_document(session: Session, doc_id: int, current_user: User) -> FileResponse:
    doc = session.get(Document, doc_id)
    if not doc:
        raise ValueError("Document not found")
        
    # Permission check
    if doc.user_id != current_user.id and not current_user.has_permission("view_all_resumes"):
        raise ValueError("You do not have permission to view this document")
        
    if not os.path.exists(doc.file_path):
        raise ValueError("File is missing from storage")
        
    # Audit log
    _log_audit(
        session=session,
        user_id=current_user.id,
        action="DOWNLOAD_DOCUMENT",
        entity_type="Document",
        entity_id=doc.id,
        details=f"Downloaded {doc.document_type} file: {doc.original_filename}"
    )
    session.commit()
    
    return FileResponse(
        path=doc.file_path,
        filename=doc.original_filename,
        media_type="application/octet-stream"
    )

def delete_document(session: Session, doc_id: int, current_user: User):
    doc = session.get(Document, doc_id)
    if not doc:
        raise ValueError("Document not found")
        
    # Permission check
    if doc.user_id != current_user.id and not current_user.has_permission("manage_users"): # Only admins can delete other's docs
        raise ValueError("You do not have permission to delete this document")
        
    # Delete physical file
    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)
        
    session.delete(doc)
    
    # Audit log
    _log_audit(
        session=session,
        user_id=current_user.id,
        action="DELETE_DOCUMENT",
        entity_type="Document",
        entity_id=doc.id,
        details=f"Deleted {doc.document_type} file: {doc.original_filename}"
    )
    session.commit()
    return True
